"""Gera documentacao tecnica em .docx com arquitetura + regras de negocio."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / 'docs' / 'Shiva_Zen_Documentacao_Tecnica.docx'


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def bullets(doc, itens):
    for i in itens:
        doc.add_paragraph(i, style='List Bullet')


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = 'Shiva Zen — Documentação Técnica — página '
    add_page_number(p)

    # ═══ Capa ═══
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SHIVA ZEN')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Sistema de Gestão de Clínica de Estética')
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run('Documentação Técnica Completa — Arquitetura, Infraestrutura e Regras de Negócio')
    r.font.size = Pt(13)
    r.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Versão 2.0 | Abril 2026\n').font.size = Pt(11)
    meta.add_run('Django 5.2.1 + Python 3.14 + PostgreSQL + Railway').font.size = Pt(11)

    doc.add_page_break()

    # ═══ Índice ═══
    doc.add_heading('Índice', level=1)
    itens = [
        '1. Visão Geral',
        '2. Arquitetura da Aplicação',
        '3. Modelo de Dados',
        '4. Autenticação e Autorização',
        '5. Comunicação Multi-Canal',
        '6. Regras de Negócio — Agendamento Público',
        '7. Regras de Negócio — Portal do Cliente',
        '8. Regras de Negócio — Portal do Profissional',
        '9. Regras de Negócio — Painel Administrativo',
        '10. Regras de Negócio — Pacotes de Sessões',
        '11. Regras de Negócio — NPS e Satisfação',
        '12. Regras de Negócio — Marketing e Promoções',
        '13. Regras de Negócio — Lista de Espera',
        '14. Regras de Negócio — Termos e Consentimentos',
        '15. LGPD e Compliance',
        '16. Segurança',
        '17. Integrações Externas',
        '18. Observabilidade',
        '19. Infraestrutura e Deploy',
        '20. API REST',
        '21. Testes',
        '22. Variáveis de Ambiente',
        '23. Limites de Escala',
        '24. E-mail Síncrono (decisão corrente)',
        '25. Roadmap de Infraestrutura',
    ]
    for i in itens:
        doc.add_paragraph(i, style='List Number')
    doc.add_page_break()

    # ═══ 1. Visão Geral ═══
    doc.add_heading('1. Visão Geral', level=1)
    doc.add_paragraph(
        'O Shiva Zen é uma plataforma web monolítica (Django 5.2.1 + Python 3.14) '
        'para gestão de uma clínica de estética de pequeno porte '
        '(aproximadamente 40 clientes/mês). Cobre o ciclo completo: agendamento '
        'público sem login, portal do cliente, portal do profissional, painel '
        'administrativo, prontuário eletrônico, pacotes de sessões, pesquisa NPS '
        'pós-atendimento, campanhas de e-mail marketing com consentimento LGPD '
        'granular e auditoria de operações sensíveis.'
    )
    doc.add_paragraph(
        'Arquitetura: Django MVT clássico, sem SPA. Templates renderizados no '
        'servidor com progressive enhancement via fetch. Service worker PWA '
        'habilitado apenas no painel administrativo. API REST v1 (DRF + '
        'drf-spectacular / OpenAPI) exposta em /api/ para futuras integrações; '
        'o fluxo principal opera com HTTP tradicional e formulários.'
    )
    doc.add_paragraph(
        'Deploy atual: Railway free tier (web + Postgres). Celery e Redis foram '
        'removidos em favor de envio de e-mail síncrono e cron externo HTTP '
        '(cron-job.org) para jobs agendados, mantendo custo em zero.'
    )

    # ═══ 2. Arquitetura ═══
    doc.add_heading('2. Arquitetura da Aplicação', level=1)
    doc.add_paragraph(
        'Código organizado em uma app Django única (app_shivazen) com separação '
        'por responsabilidade:'
    )
    bullets(doc, [
        'models/: 39 modelos em 11 arquivos por domínio (acesso, agendamentos, '
        'clientes, pacotes, procedimentos, profissionais, prontuário, sistema, '
        'nps, termos, mixins).',
        'views/: 18 módulos separados por contexto (público, booking, dashboard, '
        'admin, admin_management, admin_promotions, admin_professional, '
        'prontuário, profissional, whatsapp, lgpd, cron, health, auth, '
        'notificações, pacotes, booking_api, __init__).',
        'services/: camada de regras de negócio (agendamento, lgpd, otp, '
        'notificacao, auditoria, pacote).',
        'utils/: helpers transversais (email, whatsapp, sms, captcha, security).',
        'forms/: Django Forms com validação por recurso.',
        'api/: DRF serializers, views e urls em /api/.',
        'tasks.py: funções @shared_task, executadas síncronas via endpoint /cron/run/.',
        'middleware.py: CSP com nonce + SecurityHeadersMiddleware.',
        'signals.py: hooks Django (auditoria, derivações de status).',
        'validators.py: validadores reutilizáveis (CPF, telefone BR, data de nascimento).',
    ])

    # ═══ 3. Modelo de Dados ═══
    doc.add_heading('3. Modelo de Dados', level=1)
    doc.add_paragraph(
        'Banco PostgreSQL em produção (Railway managed). 39 modelos agrupados '
        'por domínio. Detalhamento:'
    )
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    h = t.rows[0].cells
    h[0].text = 'Domínio'
    h[1].text = 'Modelos principais'
    h[2].text = 'Responsabilidade'
    rows = [
        ('Acesso',
         'Usuario, Perfil, Funcionalidade, PerfilFuncionalidade, OTPCode',
         'Autenticação customizada (USERNAME_FIELD=email) com ACL por perfil. '
         'OTPCode armazena códigos com TTL, tentativas, canal (SMS/EMAIL), IP.'),
        ('Clientes',
         'Cliente, PreferenciaComunicacao',
         'Dados pessoais + consentimentos granulares: aceita_comunicacao '
         '(legado), consent_email_marketing, consent_whatsapp_nps, '
         'consent_whatsapp_confirmacao — cada um com timestamp e IP de origem '
         'para trilha de auditoria LGPD. Token de unsubscribe único.'),
        ('Agendamentos',
         'Atendimento, Notificacao, Bloqueio, ListaEspera, Feriado',
         'Core. Atendimento tem 7 status: PENDENTE, AGENDADO, CONFIRMADO, '
         'REALIZADO, CANCELADO, FALTOU, REAGENDADO. Notificacao rastreia '
         'envios por canal (WHATSAPP/SMS/EMAIL) e tipo.'),
        ('Procedimentos',
         'Procedimento, CategoriaProcedimento, ProfissionalProcedimento, Promocao',
         'Catálogo de serviços, categorias, relação M2M com profissionais. '
         'Slug e imagem para página pública.'),
        ('Pacotes',
         'Pacote, ItemPacote, PacoteCliente, SessaoPacote',
         'Venda de pacotes de sessões. ItemPacote define quantidade por '
         'procedimento. SessaoPacote registra consumo vinculado a atendimento.'),
        ('Profissionais',
         'Profissional, JornadaTrabalho, AprovacaoProfissional',
         'Cadastro completo, disponibilidade por dia/semana, fluxo de '
         'aprovação pelo admin antes de ativação.'),
        ('Prontuário',
         'Prontuario, AnotacaoSessao, Consentimento',
         'Registro clínico por cliente, anotações por sessão, termos '
         'assinados digitalmente.'),
        ('NPS',
         'AvaliacaoNPS, RespostaNPS, TokenNPS',
         'Pesquisa de satisfação pós-atendimento com token único de 7 dias.'),
        ('Sistema',
         'Notificacao, LogAuditoria, TermoUso, Parametro, ConfiguracaoSistema',
         'Auditoria imutável, notificações enviadas, termos versionados, '
         'parâmetros de configuração do sistema.'),
    ]
    for dom, mod, resp in rows:
        r = t.add_row().cells
        r[0].text = dom
        r[1].text = mod
        r[2].text = resp

    doc.add_paragraph(
        'Soft-delete é aplicado a Cliente via campo deletado_em. Default '
        'manager (ClienteAtivosManager) oculta registros deletados; all_objects '
        'expõe histórico completo para auditoria.'
    )
    doc.add_paragraph(
        'Constraints notáveis em Atendimento: status ∈ conjunto fixo '
        '(CheckConstraint), data_hora_fim > data_hora_inicio. '
        'Índices compostos em (cliente, status), (profissional, data_hora_inicio), '
        '(tipo, canal, status_envio) otimizam queries críticas.'
    )

    # ═══ 4. Autenticação ═══
    doc.add_heading('4. Autenticação e Autorização', level=1)
    doc.add_paragraph('Três perfis distintos com fluxos de autenticação separados:')
    bullets(doc, [
        'Cliente anônimo: acessa agendamento público sem login. Autenticação '
        'via OTP de 6 dígitos entregue por SMS (Zenvia), validade 10 minutos, '
        'máximo 5 tentativas. Gravada em OTPCode (email + hash do código + '
        'canal + telefone + IP). Renvio bloqueado por REENVIO_MINIMO_SEG.',
        'Profissional: login com senha (hash pbkdf2_sha256). Acessa apenas a '
        'própria agenda, marca atendimentos como REALIZADO, registra anotações '
        'de sessão, aprova pedidos pendentes.',
        'Admin: login com senha + django-axes (5 tentativas falhas → lockout '
        '1h por combinação IP + username). Acesso total ao painel.',
    ])
    doc.add_paragraph(
        'AuthenticationBackends: AxesStandaloneBackend (primeiro — bloqueia '
        'antes de validar credencial) + ModelBackend padrão. Lockout cobre '
        'tanto tentativas válidas quanto inválidas.'
    )

    # ═══ 5. Comunicação ═══
    doc.add_heading('5. Comunicação Multi-Canal', level=1)
    doc.add_paragraph(
        'Regra corporativa vigente, centralizada em services/notificacao.py '
        'nas classes OTPService, EmailService e WhatsAppService:'
    )
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    h = t.rows[0].cells
    h[0].text = 'Canal'
    h[1].text = 'Uso exclusivo'
    h[2].text = 'Consentimento exigido'
    for canal, uso, consent in [
        ('SMS (Zenvia)',
         'OTP de autenticação para agendamento e portal.',
         'Implícito (cliente fornece telefone = aceita receber o código).'),
        ('E-mail (SMTP)',
         'Promoções, campanhas de marketing, informações sobre pacotes de '
         'sessões, confirmações formais de compra, aniversário.',
         'consent_email_marketing para marketing/promoções/aniversário; '
         'transacional (pacote expirando/compra) não exige consent adicional.'),
        ('WhatsApp (Meta)',
         'Lembrete D-1 (24h antes do atendimento) e pesquisa NPS '
         'pós-atendimento.',
         'consent_whatsapp_confirmacao para D-1; consent_whatsapp_nps para NPS.'),
    ]:
        r = t.add_row().cells
        r[0].text = canal
        r[1].text = uso
        r[2].text = consent
    doc.add_paragraph(
        'Cada disparo verifica o consentimento antes de chamar o provedor. '
        'Tentativas bloqueadas por falta de consent são registradas em log '
        'com nível INFO (não disparam erro).'
    )
    doc.add_paragraph('Endpoints de substituição ao Celery Beat (cron externo):')
    bullets(doc, [
        'POST /cron/run/lembrete_diario/ — envia lembretes D-1 via WhatsApp.',
        'POST /cron/run/nps_24h/ — envia pesquisa NPS 24h após REALIZADO.',
        'POST /cron/run/detrator_alerta/ — alerta admin sobre nota ≤ 6.',
        'POST /cron/run/pacote_expirando/ — aviso de pacote próximo do vencimento.',
        'POST /cron/run/pacote_expirar/ — marca pacotes vencidos como EXPIRADO.',
        'POST /cron/run/aniversario/ — e-mail de aniversário.',
        'POST /cron/run/limpeza_status/ — marca FALTOU em atendimentos não atualizados.',
        'POST /cron/run/lgpd_purgar/ — anonimiza clientes inativos (2 anos).',
    ])
    doc.add_paragraph(
        'Autenticação: header X-Cron-Token comparado com hmac.compare_digest '
        'contra env var CRON_TOKEN.'
    )

    # ═══ 6. Agendamento Público ═══
    doc.add_heading('6. Regras de Negócio — Agendamento Público', level=1)
    doc.add_paragraph(
        'Fluxo de 6 etapas renderizadas na mesma página (booking-form) via JS '
        'controlando step-panels. URL: /agendamento/. View principal: '
        'views.booking.agendamento_publico + confirmar_agendamento.'
    )
    for i, etapa in enumerate([
        'Seleção de procedimento: grid filtrável por categoria. Só procedimentos '
        'com disponibilidade configurada aparecem.',
        'Seleção de profissional: opcional (default: qualquer disponível). Se '
        'selecionado, restringe slots aos profissionais do procedimento.',
        'Seleção de data/horário: API GET /api/dias-disponiveis/ retorna '
        'calendário; GET /api/horarios-disponiveis/ retorna slots livres. '
        'Considera JornadaTrabalho do profissional, bloqueios, feriados BR '
        '(carregados por management command carregar_feriados) e duração '
        'do procedimento. Não mostra horários no passado.',
        'Dados do cliente: nome, e-mail obrigatório, telefone obrigatório '
        '(exigido para OTP), data de nascimento opcional, checkboxes LGPD '
        '(email marketing, WhatsApp D-1 default marcado, WhatsApp NPS).',
        'OTP: POST /agendamento/otp/solicitar/ dispara SMS Zenvia (rate limit '
        '5/min por IP). Código 6 dígitos, validade 10min, máximo 5 tentativas. '
        'Sem fallback e-mail — se telefone inválido, fluxo aborta. POST '
        '/agendamento/otp/verificar/ valida (rate limit 10/min por IP). '
        'Aprovação do OTP persiste na sessão (otp_agendamento_email).',
        'Confirmação: POST /agendamento/confirmar/ dentro de transaction.atomic '
        'com SELECT FOR UPDATE no slot. Cloudflare Turnstile valida o token '
        'se TURNSTILE_ENABLED. Checa colisão de horário via '
        'AgendamentoService.criar (existe-conflito entre data_hora_inicio e '
        'data_hora_fim no status AGENDADO/PENDENTE/CONFIRMADO do profissional). '
        'Cria cliente se inexistente (select_for_update + get_or_create) ou '
        'atualiza campos pendentes. Registra Atendimento com status PENDENTE. '
        'Grava consents com timestamp e IP. Gera token_cancelamento único.',
    ], start=1):
        doc.add_paragraph(f'{i}. {etapa}')

    doc.add_paragraph('Regras críticas:')
    bullets(doc, [
        'OTP requer telefone — sem telefone, solicitação retorna '
        'telefone_ausente. Para cliente existente, telefone cadastrado tem '
        'prioridade sobre o digitado no form.',
        'Rate limits: 5 solicitações de OTP/minuto por IP; 10 verificações/minuto; '
        '3 SMS/hora por telefone; 10 SMS/hora por IP; 60 SMS/hora global.',
        'Cliente com bloqueado_online=True não consegue agendar online. '
        'Bloqueio automático após 3 faltas consecutivas.',
        'Após o confirmar, tela de sucesso (agendamento_sucesso.html) mostra '
        'link para "Meus Agendamentos" — não há e-mail de confirmação.',
        'Lembrete D-1 por WhatsApp (se consent) sai automaticamente via cron '
        'às 08:00 no dia anterior, para atendimentos com status AGENDADO.',
    ])

    # ═══ 7. Portal do Cliente ═══
    doc.add_heading('7. Regras de Negócio — Portal do Cliente', level=1)
    doc.add_paragraph('URL: /meus-agendamentos/. Views: booking.meus_agendamentos + fluxo OTP próprio.')
    bullets(doc, [
        'Acesso via OTP por SMS (não requer senha). Cliente informa e-mail, '
        'sistema envia código ao telefone cadastrado, cliente valida.',
        'Lista atendimentos futuros e histórico recente com '
        'AgendamentoService.ativos_no_futuro_do_cliente (select_related + '
        'order_by data_hora_inicio).',
        'Cancelamento: POST /meus-agendamentos/cancelar/<id>/ — só permitido '
        'se status ∈ {PENDENTE, AGENDADO, CONFIRMADO} e data futura. '
        'Atualiza status para CANCELADO, registra auditoria.',
        'Reagendamento: /reagendar/<token>/ usa token_cancelamento do '
        'Atendimento. View GET mostra calendário; POST cria novo Atendimento '
        'com reagendado_de=atendimento_original e marca original como REAGENDADO.',
        'Logout via POST /meus-agendamentos/logout/ limpa sessão.',
        'Confirmação de presença: link em e-mail D-1 (se canal ativo). '
        'Endpoint /confirmar-presenca/<token>/ aceita confirmar ou cancelar.',
    ])

    # ═══ 8. Portal do Profissional ═══
    doc.add_heading('8. Regras de Negócio — Portal do Profissional', level=1)
    doc.add_paragraph('URL: /profissional/. Requer autenticação com senha.')
    bullets(doc, [
        'Agenda diária/semanal visível apenas para o profissional logado.',
        'Aprovar/rejeitar atendimento com status PENDENTE: muda para AGENDADO '
        'ou CANCELADO. Dispara e-mail de confirmação ao cliente '
        '(enviar_confirmacao_agendamento_email) e alerta admin se rejeição.',
        'Marcar como REALIZADO: POST /profissional/atendimento/<id>/realizado/. '
        'Dispara PacoteService.debitar_sessao_por_atendimento (se aplicável). '
        'Permite entrada de valor_cobrado diferente do tabelado.',
        'Anotações de sessão: formulário próprio, salva em AnotacaoSessao '
        'vinculado ao atendimento. Dados sensíveis — registrados em '
        'LogAuditoria.',
        'Nenhuma ação retroativa em atendimentos REALIZADOS (exceto anotação).',
    ])

    # ═══ 9. Painel Admin ═══
    doc.add_heading('9. Regras de Negócio — Painel Administrativo', level=1)
    doc.add_paragraph('URL: /painel/. Requer autenticação com perfil Administrador.')
    bullets(doc, [
        'Overview: dashboard com KPIs (agendamentos do dia, faturamento '
        'mensal, taxa de comparecimento, NPS médio).',
        'Agendamentos: CRUD completo. Admin pode criar/editar/cancelar '
        'qualquer atendimento, mesmo no passado.',
        'Clientes: lista + detalhe. Ver consentimentos ativos, histórico de '
        'atendimentos, avaliações NPS. Bloquear/desbloquear online.',
        'Profissionais: cadastro + edição. Define jornada (dias/horários '
        'de trabalho), procedimentos que atende, status (ATIVO/INATIVO).',
        'Procedimentos: CRUD + categoria. Define duração, preço, imagem.',
        'Pacotes: CRUD + venda. Composto por itens (procedimento + quantidade '
        'de sessões). Admin registra venda em PacoteCliente.',
        'Promoções: CRUD + disparo de e-mail. Campanha valida '
        'consent_email_marketing por destinatário.',
        'Bloqueios de agenda: admin registra período de indisponibilidade '
        '(férias, evento) por profissional.',
        'Lista de espera: visualiza inscritos, notifica quando houver '
        'disponibilidade.',
        'Auditoria: /painel/auditoria/ consulta LogAuditoria com filtros por '
        'operação, usuário, data.',
        'Termos: gerencia versões de TermoUso; exibe assinaturas pendentes.',
        'Prontuário: acessa registros clínicos com auditoria obrigatória '
        '(Consentimento explícito antes de ver dados sensíveis).',
        'Relatórios: exporta Excel (xlsx) com agendamentos/faturamento por '
        'período. View exportar_relatorio_excel.',
        'Preview de e-mail: renderiza template com dados de exemplo antes '
        'de disparar campanha.',
    ])

    # ═══ 10. Pacotes ═══
    doc.add_heading('10. Regras de Negócio — Pacotes de Sessões', level=1)
    doc.add_paragraph(
        'Pacote é uma coleção de ItemPacote (procedimento + quantidade). '
        'Cliente compra via PacoteCliente, que tem status ATIVO/EXPIRADO/FINALIZADO.'
    )
    bullets(doc, [
        'Compra: admin registra PacoteCliente.criar (status=ATIVO, '
        'data_expiracao = data_compra + validade_dias do Pacote).',
        'Consumo: quando atendimento é marcado como REALIZADO, '
        'PacoteService.debitar_sessao_por_atendimento procura o primeiro '
        'PacoteCliente ATIVO do cliente que tenha ItemPacote cobrindo o '
        'procedimento e ainda possua saldo de sessões (sessoes_ja_feitas < '
        'quantidade_sessoes). Cria SessaoPacote.',
        'Finalização: se todas as sessões foram consumidas, '
        'PacoteCliente.verificar_finalizacao marca status=FINALIZADO.',
        'Expiração: job pacote_expirar (cron diário 00:30) percorre '
        'PacoteCliente ATIVOS com data_expiracao < hoje e marca '
        'status=EXPIRADO.',
        'Aviso prévio: job pacote_expirando (cron diário 07:00) envia '
        'e-mail 7 dias antes da expiração (transacional, não exige consent).',
        'Ordenação: pacote mais antigo é consumido primeiro (order_by '
        'criado_em).',
    ])

    # ═══ 11. NPS ═══
    doc.add_heading('11. Regras de Negócio — NPS e Satisfação', level=1)
    bullets(doc, [
        'Disparo: job nps_24h (cron diário 10:00) seleciona atendimentos '
        'com status=REALIZADO, data_hora_fim ≤ agora-1dia, sem '
        'AvaliacaoNPS vinculada, com consent_whatsapp_nps=True.',
        'Canal: WhatsApp Meta exclusivo. Gera token (secrets.token_urlsafe(32)) '
        'e link /nps/<token>/. Mensagem via template Meta aprovado.',
        'Expiração do token: 7 dias após envio. Link expirado mostra aviso.',
        'Resposta: página pública mostra escala 0-10 + campo comentário. '
        'Salva AvaliacaoNPS com nota e comentário.',
        'Categorização: nota 0-6 = detrator; 7-8 = neutro; 9-10 = promotor.',
        'Alerta admin detrator: job detrator_alerta (cron diário 10:30) '
        'envia e-mail ao email_admin para cada avaliação detratora com '
        'alerta_enviado=False, depois marca como enviado (idempotência).',
        'Sem consent: NPS é ignorado para o cliente; não há fallback via '
        'e-mail por padrão.',
    ])

    # ═══ 12. Marketing ═══
    doc.add_heading('12. Regras de Negócio — Marketing e Promoções', level=1)
    bullets(doc, [
        'Cadastro: admin cria Promocao (título, corpo HTML, cupom opcional, '
        'validade em dias). Preview disponível.',
        'Disparo: admin clica em "Enviar" — sistema percorre Clientes com '
        'consent_email_marketing=True e e-mail preenchido. Para cada um, '
        'chama EmailService.enviar_promocao passando o token de unsubscribe.',
        'Campanha registrada em Notificacao (canal=EMAIL, tipo=CONFIRMACAO '
        'ou criado específico para promoção).',
        'Opt-out one-click: cabeçalho RFC 8058 List-Unsubscribe + link no '
        'footer do e-mail. GET /lgpd/unsubscribe/<token>/ marca '
        'consent_email_marketing=False, consent_whatsapp_nps=False e '
        'aceita_comunicacao=False.',
        'Aniversário: job aniversario (cron diário 09:00) percorre clientes '
        'com aniversário no dia + consent_email_marketing, envia '
        'EmailService.enviar_aniversario.',
    ])

    # ═══ 13. Lista de Espera ═══
    doc.add_heading('13. Regras de Negócio — Lista de Espera', level=1)
    bullets(doc, [
        'Inscrição: /lista-espera/ formulário público. Cliente informa nome, '
        'e-mail, telefone, procedimento de interesse, data-alvo aproximada.',
        'Cria ListaEspera com status AGUARDANDO.',
        'Notificação: quando admin cria um bloqueio que libera horários, ou '
        'quando há cancelamento, admin pode acionar notificação manual via '
        'painel. Task job_notificar_fila_espera (sem disparo agendado) envia '
        'WhatsApp template convidando a agendar.',
        'Sem disparo automático — evita spam e controla fluxo manualmente.',
    ])

    # ═══ 14. Termos ═══
    doc.add_heading('14. Regras de Negócio — Termos e Consentimentos', level=1)
    bullets(doc, [
        'TermoUso versionado: admin cria nova versão quando há alteração de '
        'texto. Versão anterior permanece imutável como evidência.',
        'Assinatura do cliente: /termo/assinar/<token>/ apresenta o texto; '
        'cliente clica em "Aceito", sistema grava Consentimento com cliente, '
        'versao_termo, data, IP.',
        'Termos pendentes: se cliente tiver atendimento mas não tiver '
        'Consentimento na versão atual, job ao criar atendimento envia '
        'e-mail enviar_termos_pendentes_email solicitando assinatura.',
        'Prontuário exige Consentimento específico (diferente de termo geral). '
        'prontuario_consentimento registra aceite explícito antes da coleta '
        'de dados clínicos.',
    ])

    # ═══ 15. LGPD ═══
    doc.add_heading('15. LGPD e Compliance', level=1)
    bullets(doc, [
        'Consentimentos granulares por canal em Cliente, cada um com timestamp '
        'e IP de origem. Captura no agendamento público.',
        'DSAR (Direito do Titular): /lgpd/meus-dados/ — autenticado via OTP, '
        'exibe JSON ou HTML com todos os dados do cliente. Backend: '
        'LgpdService.exportar_dados_cliente (dados pessoais, preferências, '
        'atendimentos, avaliações NPS).',
        'Direito ao esquecimento: LgpdService.esquecer_cliente anonimiza '
        'campos pessoais (nome → [ANONIMIZADO-<pk>], demais → None) e '
        'soft-delete.',
        'Opt-out one-click: /lgpd/unsubscribe/<token>/ zera todos consents '
        '(aceita_comunicacao, consent_email_marketing, consent_whatsapp_nps).',
        'Retenção: job lgpd_purgar (cron semanal domingo 03:00) aplica '
        'esquecer_cliente em clientes sem atendimento há mais de 2 anos '
        '(RETENCAO_CLIENTE_INATIVO_DIAS=730).',
        'Log auditoria: LogAuditoria grava toda operação sensível (acesso '
        'prontuário, exportação relatório, alteração de consent, login admin). '
        'Retenção 1 ano.',
        'DPIA: documento docs/DPIA.md registra impacto, riscos e mitigações.',
    ])

    # ═══ 16. Segurança ═══
    doc.add_heading('16. Segurança', level=1)
    bullets(doc, [
        'CSP com nonce por request (ContentSecurityPolicyMiddleware). Bloqueia '
        'inline scripts não autorizados.',
        'SecurityHeadersMiddleware: HSTS (1 ano + preload), X-Frame-Options '
        'DENY, X-Content-Type-Options nosniff, Referrer-Policy '
        'strict-origin-when-cross-origin, Cross-Origin-Opener-Policy same-origin.',
        'CSRF: tokens em todos os formulários. SAMESITE=Lax. '
        'CSRF_COOKIE_SECURE em produção.',
        'HTTPS forçado via SECURE_SSL_REDIRECT. Healthchecks (/healthz/, '
        '/health/) excluídos do redirect para probes HTTP internos do Railway.',
        'django-axes: lockout de brute-force (5 tentativas por '
        'combinação IP+username, timeout 1h).',
        'Rate limits defense-in-depth: OTP 5/min IP; verificação OTP 10/min; '
        'SMS 3/hora telefone + 10/hora IP + 60/hora global.',
        'Zenvia webhook: IP allowlist + HMAC do payload via '
        'ZENVIA_WEBHOOK_SECRET.',
        'Cron endpoint: X-Cron-Token validado com hmac.compare_digest.',
        'Senhas: pbkdf2_sha256 padrão Django. Mínimo 10 caracteres '
        '(MinimumLengthValidator).',
        'PII masking em logs: LogAuditoria grava hash de e-mail/telefone, '
        'nunca valor em texto puro.',
        'Secrets via env var, nunca commitados. .env.example documenta.',
    ])

    # ═══ 17. Integrações ═══
    doc.add_heading('17. Integrações Externas', level=1)
    bullets(doc, [
        'SMTP (Gmail App Password ou SendGrid): envio transacional e '
        'marketing por e-mail.',
        'Meta WhatsApp Business API: templates aprovados para lembrete D-1 '
        'e NPS. Webhook /api/whatsapp/webhook/ para recebimento de status.',
        'Zenvia SMS API v2: envio de OTP. Webhook /api/zenvia/webhook/ '
        'para delivery reports. Rate limit + IP allowlist + HMAC.',
        'Cloudflare Turnstile: CAPTCHA invisível em OTP e confirmação de '
        'agendamento. Opcional (ativado via TURNSTILE_SITE_KEY).',
        'Sentry (opcional): error reporting e tracing. Ativado quando '
        'SENTRY_DSN definido. send_default_pii=False.',
        'cron-job.org (externo, gratuito): chama POST /cron/run/<job>/ nos '
        'horários definidos, com header X-Cron-Token.',
    ])

    # ═══ 18. Observabilidade ═══
    doc.add_heading('18. Observabilidade', level=1)
    bullets(doc, [
        '/healthz/: liveness — retorna 200 com {"status": "alive"} se processo '
        'vivo. Usado pelo healthcheck do Railway.',
        '/health/: readiness — checa DB + cache. Retorna 503 em degradação. '
        'Parâmetro ?celery=1 força check adicional de broker (se configurado).',
        'Logging JSON estruturado em produção (python-json-logger). Nível '
        'WARNING root, INFO django, INFO app_shivazen.',
        'Sentry integra Django via DjangoIntegration. TRACES_SAMPLE_RATE '
        'configurável (default 0.2).',
        'LogAuditoria (DB): trilha imutável de operações sensíveis; admin '
        'consulta em /painel/auditoria/ com filtros.',
        'Railway logs V2 (JSON): console do serviço acessível via CLI ou '
        'dashboard. Voláteis — sem retenção longa no free tier.',
    ])

    # ═══ 19. Infra ═══
    doc.add_heading('19. Infraestrutura e Deploy', level=1)
    doc.add_paragraph('Configuração corrente no Railway free tier:')
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    h = t.rows[0].cells
    h[0].text = 'Serviço'
    h[1].text = 'Configuração'
    h[2].text = 'Custo'
    for srv, cfg, custo in [
        ('web',
         'Django + gunicorn (1 worker / 4 threads / timeout 120s). Nixpacks '
         'build. sleep após ~5min sem tráfego. Domínio '
         'web-production-465af.up.railway.app.',
         'Free'),
        ('Postgres',
         'Managed Railway. 1GB volume. DATABASE_URL injetada via '
         'referência de serviço.',
         'Free'),
        ('Cron externo',
         'cron-job.org chama endpoints HTTP com X-Cron-Token.',
         'Free'),
    ]:
        r = t.add_row().cells
        r[0].text = srv
        r[1].text = cfg
        r[2].text = custo
    doc.add_paragraph('Ambientes Railway (mesmos serviços, vars separadas):')
    bullets(doc, [
        'production: deploy automático do branch main via Railway GitHub '
        'integration. Healthcheck em /healthz/.',
        'dev: deploy automático do branch dev via deployment trigger '
        'GraphQL. Usado para testes antes do merge em main.',
    ])
    doc.add_paragraph('Configuração de build (railway.json):')
    bullets(doc, [
        'builder: NIXPACKS (detecção automática Django + requirements.txt).',
        'startCommand: sh -c "python manage.py migrate --noinput && '
        'python manage.py collectstatic --noinput --verbosity 0 && exec '
        'gunicorn shivazen.wsgi --bind 0.0.0.0:${PORT:-8080} --workers 1 '
        '--threads 4 --timeout 120 --access-logfile - --error-logfile -".',
        'healthcheckPath: /healthz/, healthcheckTimeout: 60s.',
        'restartPolicyType: ON_FAILURE, maxRetries: 10.',
    ])

    # ═══ 20. API ═══
    doc.add_heading('20. API REST', level=1)
    doc.add_paragraph(
        'API REST v1 em /api/ com drf-spectacular (OpenAPI 3.0). Schema em '
        '/api/schema/, Swagger UI em /api/docs/.'
    )
    bullets(doc, [
        'Autenticação: SessionAuthentication (cookie). Endpoints públicos '
        '(ex. listagem de procedimentos) usam AnonRateThrottle.',
        'Throttling: anon 60/hora, user 1000/hora.',
        'Paginação: PageNumberPagination, PAGE_SIZE 20.',
        'Endpoints principais: GET /api/procedimentos/, GET '
        '/api/horarios-disponiveis/, GET /api/dias-disponiveis/.',
        'Boas práticas DRF: serializers separados para leitura/escrita, '
        'componentização via COMPONENT_SPLIT_REQUEST.',
    ])

    # ═══ 21. Testes ═══
    doc.add_heading('21. Testes', level=1)
    doc.add_paragraph(
        '135 testes unitários + integração em app_shivazen/tests/. SQLite em '
        'memória (override em DATABASES quando sys.argv contém test ou env '
        'PYTEST_CURRENT_TEST setada).'
    )
    bullets(doc, [
        'test_integration_booking: fluxo completo agendamento com OTP.',
        'test_otp: geração, expiração, rate limit.',
        'test_integration_nps: envio de token, resposta, idempotência.',
        'test_integration_packages: compra, consumo, expiração, finalização.',
        'test_whatsapp_webhook: HMAC, templates.',
        'test_security_utils: CSRF, IP extraction, LGPD services.',
        'test_validators: CPF, telefone BR, data nascimento.',
        'test_context_processors e test_decorators: helpers transversais.',
        'test_public_views: renderização de páginas públicas.',
        'test_signals: hooks automáticos (auditoria, derivação status).',
        'test_confirmar_agendamento: endpoint pós-lembrete.',
        'test_gaps_coverage: cobertura adicional de edge-cases.',
    ])

    # ═══ 22. Env Vars ═══
    doc.add_heading('22. Variáveis de Ambiente', level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    h = t.rows[0].cells
    h[0].text = 'Variável'
    h[1].text = 'Obrigatória?'
    h[2].text = 'Descrição'
    for var, req, desc in [
        ('DJANGO_SECRET_KEY', 'Sim', 'Chave de assinatura. Mínimo 50 caracteres aleatórios.'),
        ('DJANGO_ENV', 'Sim', 'prod ou dev — seleciona settings module.'),
        ('DATABASE_URL', 'Sim', 'postgres://... injetado pelo Railway.'),
        ('ALLOWED_HOSTS', 'Sim', 'Domínios separados por vírgula.'),
        ('CSRF_TRUSTED_ORIGINS', 'Sim', 'Origens com esquema https://.'),
        ('USE_HTTPS', 'Não', 'Default True. False desativa SSL redirect.'),
        ('CRON_TOKEN', 'Sim', 'Token do endpoint /cron/run/ (32+ chars aleatórios).'),
        ('EMAIL_HOST', 'Sim', 'SMTP host (ex. smtp.gmail.com).'),
        ('EMAIL_HOST_USER', 'Sim', 'Usuário SMTP.'),
        ('EMAIL_HOST_PASSWORD', 'Sim', 'Senha/App Password SMTP.'),
        ('EMAIL_PORT', 'Não', 'Default 587.'),
        ('DEFAULT_FROM_EMAIL', 'Sim', 'E-mail remetente.'),
        ('ZENVIA_API_TOKEN', 'Sim (para OTP)', 'Token Zenvia v2.'),
        ('ZENVIA_FROM', 'Sim (para OTP)', 'Identificador remetente SMS.'),
        ('ZENVIA_WEBHOOK_SECRET', 'Não', 'HMAC do webhook Zenvia.'),
        ('TURNSTILE_SITE_KEY', 'Não', 'CAPTCHA Cloudflare público.'),
        ('TURNSTILE_SECRET_KEY', 'Não', 'CAPTCHA Cloudflare server.'),
        ('SENTRY_DSN', 'Não', 'Observabilidade — ativa Sentry quando presente.'),
        ('SENTRY_TRACES_SAMPLE_RATE', 'Não', 'Default 0.2.'),
        ('AXES_FAILURE_LIMIT', 'Não', 'Default 5 tentativas.'),
        ('AXES_COOLOFF_TIME_HOURS', 'Não', 'Default 1 hora.'),
        ('SMS_MAX_POR_HORA', 'Não', 'Default 3 SMS/hora por telefone.'),
        ('SMS_MAX_POR_IP_HORA', 'Não', 'Default 10 SMS/hora por IP.'),
        ('SMS_MAX_GLOBAL_HORA', 'Não', 'Default 60 SMS/hora global.'),
        ('SITE_URL', 'Não', 'URL base (usada em links de NPS).'),
        ('CLINIC_NAME', 'Não', 'Default "Shiva Zen".'),
    ]:
        r = t.add_row().cells
        r[0].text = var
        r[1].text = req
        r[2].text = desc

    # ═══ 23. Limites ═══
    doc.add_heading('23. Limites de Escala', level=1)
    bullets(doc, [
        'Volume alvo: até 40 clientes/mês, ~50 agendamentos/mês, ~200 e-mails/mês.',
        'Serviço web dorme após ~5min sem tráfego — cold start 10-30s no '
        'primeiro request.',
        'SMTP síncrono limita request a ~3s.',
        'Sem fila de retry para e-mails: SMTP falhando faz o cliente ver erro.',
        'Postgres Railway free: 1GB. Suficiente para ~10 mil clientes.',
        'Cache LocMemCache por worker — não compartilha entre workers; '
        'irrelevante com 1 worker gunicorn.',
        'Jobs agendados dependem de cron-job.org (SLA histórico ~99.9%).',
    ])

    # ═══ 24. E-mail Síncrono ═══
    doc.add_heading('24. E-mail Síncrono (decisão corrente)', level=1)
    doc.add_paragraph(
        'O sistema envia e-mails dentro do request HTTP do Django. Latência '
        'típica 2-3 segundos. Mitigação UX: spinner global em forms com '
        'atributo data-loading-msg, exibindo "Enviando confirmação por '
        'e-mail..." durante o submit. Forms marcados:'
    )
    bullets(doc, [
        'agendamento_publico.html (confirmação de agendamento)',
        'reagendar.html (reagendamento)',
        'confirmar_presenca.html (confirmar/cancelar)',
        'contato.html (contato público)',
        'lista_espera.html (inscrição)',
    ])
    doc.add_paragraph(
        'Critério para migrar a assíncrono (Celery + Redis + worker): '
        'volume > 200 clientes/mês, SMTP > 3s ou campanhas em batch. '
        'Documento completo em docs/EMAIL_ASYNC_VS_SYNC.md.'
    )

    # ═══ 25. Roadmap ═══
    doc.add_heading('25. Roadmap de Infraestrutura', level=1)
    doc.add_paragraph('Próximos passos recomendados, em ordem de prioridade:')
    for item in [
        '1. Backup automatizado do Postgres (dump diário + cópia off-Railway) — crítico LGPD.',
        '2. Sentry: free tier cobre o volume atual.',
        '3. Cloudflare na frente (free tier): WAF + DDoS + cache de estáticos.',
        '4. Storage S3/R2 para MEDIA_ROOT (fotos de prontuário, avatares).',
        '5. Uptime monitor externo (BetterStack ou UptimeRobot).',
        '6. Log aggregation externa (BetterStack Logs): logs do Railway são '
        'voláteis; LGPD exige retenção de auditoria.',
        '7. Staging dev isolado com Postgres próprio (~$5/mês adicional).',
        '8. Redis + Celery quando volume > 200 clientes/mês (+$10-15/mês).',
        '9. Custom domain com HTTPS (Railway emite certificado automático).',
        '10. Split do app_shivazen em apps menores se crescer para > 50 modelos.',
    ]:
        doc.add_paragraph(item)

    doc.add_paragraph()
    end = doc.add_paragraph('— Fim do documento —', style='Intense Quote')
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f'Gerado: {OUT}')


if __name__ == '__main__':
    build()
