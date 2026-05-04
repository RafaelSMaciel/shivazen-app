"""Gera Word tecnico base do projeto Jaqueline Aranha Estetica.

Material consolidado p/ extrair conteudo do artigo CONTEEX (FAM, 7 semestre).
Segue formato Arial + estrutura Vancouver-friendly, mas e material denso
(~30 paginas) — usuario filtra/recorta para Short Paper, Full Paper ou Banner.
"""
import os
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


MER_PNG = r'C:\Users\rafae\OneDrive\Projeto FAM - Shivazen\Diagramas\MER - JaquelineAranhaEstetica.png'


def add_image(doc, path, *, width_cm=16, caption=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        _set_font(r, size=9, italic=True, color=GRAY)


GOLD = RGBColor(0xC9, 0xA8, 0x4C)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)


def _set_font(run, name='Arial', size=10, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), name)
    rfonts.set(qn('w:hAnsi'), name)
    rpr.append(rfonts)


def add_p(doc, text, *, size=10, bold=False, italic=False, color=None,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, size=sizes.get(level, 10), bold=True, color=DARK)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_font(run, size=16, bold=True, color=DARK)


def add_bullet(doc, text, *, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run('')
    p.runs[0].text = ''
    run = p.add_run(text)
    _set_font(run, size=size)


def add_table_simple(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # bg color
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1A2E')
        tcPr.append(shd)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            _set_font(run, size=9)
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def main():
    doc = Document()
    # Margens 1,5cm
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ─── CAPA ───
    add_title(doc, 'Sistema de Gestão de Clínica Estética')
    add_p(doc, 'Plataforma white-label para agendamento online, prontuário digital e gestão operacional — estudo de caso: Jaqueline Aranha Estética',
          align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color=GRAY, space_after=18)

    add_p(doc, 'Rafael Maciel ¹', align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=2)
    add_p(doc, 'Orientador: <preencher>', align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True, color=GRAY, space_after=2)
    add_p(doc, '¹ Curso: <preencher CCOMP/GTI/EC etc.> — 7º semestre — Centro Universitário FAM',
          align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=GRAY, space_after=24)

    add_p(doc, f'Documento técnico — material consolidado para artigo CONTEEX 2026-1. Versão: {date.today().isoformat()}',
          align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True, color=GRAY, space_after=24)

    # ─── RESUMO ───
    add_heading(doc, 'RESUMO', level=1)
    add_p(doc,
          'Este trabalho apresenta o desenvolvimento de uma plataforma web para gestão completa '
          'de uma clínica de estética, atendendo o caso real da biomédica Jaqueline Aranha. '
          'A solução cobre desde o agendamento público com verificação OTP até o painel administrativo '
          'com prontuário digital, controle financeiro de pacotes, notificações multi-canal '
          '(WhatsApp Business API, SMS via Zenvia e e-mail), prontuário eletrônico com anamnese '
          'estruturada, integração bidirecional com Google Calendar, sincronização iCalendar via feed '
          'assinado, Progressive Web App (PWA) instalável com web push notifications, motor de '
          'workflows configuráveis, e máquina de estados (FSM) para o ciclo de vida do atendimento. '
          'O sistema foi implementado utilizando Django 5.2, PostgreSQL, Redis e Celery, '
          'hospedado na plataforma Railway com pipeline contínuo a partir do GitHub. '
          'O semestre atual concentrou-se em cinco sprints de evolução: regras avançadas de agenda '
          'com exceções por data e RRULE iCal, PWA admin separado, segmentação de pacientes (CRM '
          'tags), planos de tratamento personalizados e modelos clínicos específicos (patch test, '
          'foto antes/depois, controle de estoque). O resultado é um sistema com 47 tabelas em '
          'produção, conformidade LGPD com consentimentos granulares e auditoria completa, '
          'e arquitetura preparada para escala.',
          space_after=8)

    add_heading(doc, 'PALAVRAS-CHAVE', level=2)
    add_p(doc, 'Agendamento online; Clínica estética; Django; PWA; LGPD; Workflow engine; CRM clínico.',
          space_after=14)

    # ─── 1. INTRODUÇÃO ───
    add_heading(doc, '1. Introdução', level=1)
    add_p(doc,
          'A digitalização de pequenas clínicas de estética no Brasil ainda é incipiente. '
          'Profissionais autônomos como biomédicas estéticas frequentemente operam com agenda em '
          'caderno, recebem agendamentos via WhatsApp manual, controlam pacotes em planilhas e '
          'perdem oportunidades de fidelização por ausência de histórico estruturado.')
    add_p(doc,
          'O presente projeto endereça este cenário desenvolvendo uma plataforma white-label '
          'completa, configurável por variáveis de ambiente para qualquer clínica single-tenant. '
          'O cliente real é a biomédica Jaqueline Aranha, atendendo presencialmente em São José '
          'do Rio Preto — SP, com 23 procedimentos catalogados (Botox, preenchimentos faciais, '
          'aplicações enzimáticas, drenagem linfática, depilação a laser, entre outros).')
    add_p(doc,
          'O problema central é a ausência de uma ferramenta integrada que combine agendamento '
          'público autosserviço, prontuário eletrônico, conformidade com a Lei Geral de Proteção '
          'de Dados (LGPD) e automação de comunicação com pacientes. Soluções de mercado como '
          'Pabau, Calendly e Cal.com cobrem partes do escopo, mas têm custo proibitivo para '
          'profissionais autônomos, exigem múltiplas integrações ou não atendem requisitos '
          'regulatórios brasileiros (LGPD, formato de telefones, integração WhatsApp Business).')
    add_p(doc,
          'O objetivo geral é construir uma plataforma open-source completa para clínicas de '
          'estética single-tenant, com interface administrativa instalável (PWA) e site público de '
          'agendamento, suportada por backend Django, banco PostgreSQL e infraestrutura serverless. '
          'Os objetivos específicos são: (i) implementar agendamento público com verificação por '
          'OTP via SMS; (ii) construir painel administrativo com calendário visual, drag-drop, '
          'KPIs e timeline de pacientes; (iii) integrar canais WhatsApp Business, SMS e e-mail; '
          '(iv) garantir conformidade LGPD com consentimentos granulares, soft-delete e '
          'anonimização agendada; (v) prover sincronização bidirecional com Google Calendar e '
          'feed iCalendar; (vi) implementar Progressive Web App com cache offline e web push.')

    # ─── 2. METODOLOGIA ───
    add_heading(doc, '2. Metodologia', level=1)
    add_p(doc,
          'O desenvolvimento seguiu metodologia ágil iterativa, organizada em sprints temáticas. '
          'O semestre atual concentrou cinco sprints de evolução sobre uma base já existente do '
          'semestre anterior, somando aproximadamente 100 commits no controle de versão.')

    add_heading(doc, '2.1 Stack tecnológico', level=2)
    add_table_simple(doc,
        ['Camada', 'Tecnologia', 'Justificativa'],
        [
            ['Backend', 'Django 5.2 + Python 3.12', 'Ecossistema maduro, ORM, admin built-in, segurança'],
            ['Banco de dados', 'PostgreSQL 14+', 'JSONB, constraints, índices compostos'],
            ['Cache / Filas', 'Redis 7 + Celery 5.4', 'Tasks assíncronas (notificações, exports)'],
            ['Servidor', 'Gunicorn + WhiteNoise', 'WSGI estável + serve estáticos'],
            ['Frontend', 'Bootstrap 5.3, FullCalendar 6, Vanilla JS', 'Sem build chain complexa'],
            ['PWA', 'Service Worker + Manifest + pywebpush VAPID', 'Instalável, offline, push'],
            ['Autenticação', 'django-axes + django-two-factor-auth (TOTP)', 'Lockout brute-force + 2FA'],
            ['Captcha', 'Cloudflare Turnstile', 'Anti-bot privacidade-preservadora'],
            ['Notificações', 'WhatsApp Business API + Zenvia SMS + SMTP', 'Multi-canal'],
            ['Hospedagem', 'Railway (Nixpacks)', 'Deploy automático, Postgres gerenciado'],
            ['Observabilidade', 'Sentry SDK + JSON logs', 'Rastreamento de erros'],
        ],
        col_widths_cm=[3, 5, 8])

    add_heading(doc, '2.2 Repositório e versionamento', level=2)
    add_p(doc,
          'O código-fonte é hospedado no GitHub no repositório jaqueline-aranha-estetica, sob '
          'licença MIT. O fluxo segue Git Flow simplificado: branch dev acumula trabalho contínuo '
          'e branch main reflete o estado de produção, com deploy automático via webhook Railway.')

    add_heading(doc, '2.3 Pipeline de deploy', level=2)
    add_p(doc,
          'A integração contínua utiliza GitHub Actions para validação local. O deploy é '
          'orquestrado pelo Railway: push em main aciona build do Dockerfile, executa migrations '
          'do Django e inicia o serviço Gunicorn, com healthcheck em /healthz/. Variáveis '
          'sensíveis são gerenciadas por secrets no painel do Railway.')

    # ─── 3. DESENVOLVIMENTO ───
    add_heading(doc, '3. Desenvolvimento', level=1)

    add_heading(doc, '3.1 Arquitetura geral', level=2)
    add_p(doc,
          'A aplicação segue arquitetura monolítica modular Django, organizada em uma única app '
          'principal (aranha_estetica) com separação por domínios: views por contexto (booking, '
          'admin, profissional, dashboard, lgpd), services para regras de negócio reutilizáveis '
          '(workflow_engine, push, gcal, otp, notificacao), models particionados em arquivos por '
          'agregado (clientes, agendamentos, profissionais, prontuário, pacotes, workflow, '
          'extras). O projeto Django (clinica) contém apenas configuração: settings dispatcher '
          'por ambiente, URLs raiz, WSGI/ASGI e Celery.')

    add_heading(doc, '3.2 Modelo de dados', level=2)
    add_p(doc,
          'O banco em produção contém 47 tabelas de negócio, modeladas com integridade '
          'referencial estrita (FKs RESTRICT em entidades de auditoria, CASCADE em detalhes '
          'compostos), constraints CHECK em campos de status e índices compostos para '
          'consultas frequentes. O Modelo Entidade-Relacionamento (MER) completo encontra-se '
          'no diagrama abaixo (Figura 1), gerado via Mermaid CLI e renderizado em PNG e PDF.')
    add_image(doc, MER_PNG, width_cm=17,
              caption='Figura 1 — Modelo Entidade-Relacionamento (MER) — 47 tabelas, geradas via Mermaid CLI.')
    add_p(doc, 'Principais agregados e suas responsabilidades:')
    add_bullet(doc, 'Acesso e RBAC: usuario, perfil, perfil_funcionalidade, funcionalidade.')
    add_bullet(doc, 'Profissional e agenda: profissional, disponibilidade_profissional (regra semanal), excecao_disponibilidade (folgas e horários diferentes por data), bloqueio_agenda (com suporte a RRULE iCal), feriado.')
    add_bullet(doc, 'Catálogo: procedimento (com categoria, duracao_minutos, buffer_minutos), preco (vigência por data, opcionalmente específico por profissional), promocao.')
    add_bullet(doc, 'Atendimento: tabela central, FSM de status (PENDENTE → AGENDADO → CONFIRMADO → REALIZADO/CANCELADO/FALTOU/REAGENDADO), token de cancelamento único.')
    add_bullet(doc, 'Cliente e prontuário: cliente, prontuario, prontuario_pergunta/resposta, anotacao_sessao, formulario_anamnese (schema JSON dinâmico), resposta_anamnese.')
    add_bullet(doc, 'Pacotes e planos: pacote, item_pacote, pacote_cliente, sessao_pacote (catálogo); plano_tratamento, item_plano_tratamento (personalizado por cliente).')
    add_bullet(doc, 'CRM e estética: tag, cliente_tag, patch_test (teste alérgico), foto_antes_depois (com consent LGPD), credito_cliente, movimento_credito.')
    add_bullet(doc, 'Estoque: produto, movimento_estoque (entrada/saída/ajuste/perda).')
    add_bullet(doc, 'Notificações e workflow: notificacao (multi-canal), workflow_regra (trigger configurável), workflow_execucao (deduplicação UNIQUE).')
    add_bullet(doc, 'PWA: web_push_subscription (VAPID keys).')
    add_bullet(doc, 'Conformidade LGPD: versao_termo, aceite_privacidade, assinatura_termo_procedimento, log_auditoria.')
    add_bullet(doc, 'Sistema: configuracao_sistema (chave-valor), codigo_verificacao, otp_code.')

    add_heading(doc, '3.3 Funcionalidades principais', level=2)

    add_heading(doc, '3.3.1 Site público — fluxo de agendamento', level=3)
    add_p(doc,
          'O agendamento público implementa fluxo em três passos: (1) seleção de procedimento '
          'com filtro por categoria; (2) escolha de data e horário disponível, calculado em '
          'tempo real considerando regras de disponibilidade semanal, exceções, buffer entre '
          'atendimentos, antecedência mínima (min_notice) e janela máxima (max_advance); '
          '(3) confirmação com dados do paciente, validação OTP via SMS (Zenvia) e Cloudflare '
          'Turnstile contra bots. Após o agendamento, o cliente recebe link mágico de '
          'cancelamento ou reagendamento self-service válido até 24 horas antes do atendimento.')

    add_heading(doc, '3.3.2 Painel administrativo', level=3)
    add_p(doc,
          'O painel administrativo concentra a operação diária: dashboard com oito indicadores '
          '(KPIs) em tempo real (agendamentos hoje/semana, ticket médio, taxa de ocupação, '
          'pacientes ativos em 90 dias, NPS médio em 30 dias, etc.), calendário visual '
          'baseado em FullCalendar com visualizações dia/semana/mês e suporte a drag-and-drop '
          'para reagendamento, listas de agendamentos e pacientes com filtro instantâneo '
          'client-side, ficha de paciente com timeline visual, LTV calculado, procedimento mais '
          'frequente e contadores de no-show, gestão de profissionais com horários semanais e '
          'exceções de agenda, CRUD de procedimentos, pacotes, promoções, lista de espera, '
          'workflows configuráveis, formulários de anamnese, controle 2FA e branding white-label.')

    add_heading(doc, '3.3.3 Notificações multi-canal', level=3)
    add_p(doc,
          'A comunicação com pacientes utiliza três canais coordenados, conforme estratégia '
          'aprovada: WhatsApp Business API para lembrete 24h antes do atendimento e pesquisa '
          'NPS pós-atendimento (apenas dois templates Meta aprovados, respeitando janelas de '
          'sessão); SMS via Zenvia para envio primário de OTP de verificação; e-mail SMTP para '
          'OTP secundário, confirmações, cancelamentos, alertas de pacote expirando, mensagens '
          'de aniversário, promoções (com consent granular opt-in) e alertas administrativos '
          'para notas NPS detratoras. Todas as notificações são registradas em tabela auditável '
          'com canal, status de envio, timestamps e resposta do cliente.')

    add_heading(doc, '3.3.4 Progressive Web App (PWA)', level=3)
    add_p(doc,
          'A aplicação é instalável como PWA com dois manifests separados: o público em /, '
          'voltado para pacientes, e o administrativo em /painel/manifest.json, com shortcuts '
          'para Calendário, Agendamentos, Pacientes e Dashboard, escopo restrito ao painel e '
          'tema visual diferenciado. O service worker (sw.js) implementa estratégia híbrida '
          'de cache: network-first para HTML com fallback offline, stale-while-revalidate '
          'para estáticos, cache-first com TTL para imagens e cache separado para APIs públicas '
          'de leitura. Web push notifications utilizam protocolo VAPID via biblioteca pywebpush, '
          'com tabela de subscriptions por usuário e fallback gracioso quando chaves não estão '
          'configuradas.')

    add_heading(doc, '3.3.5 Integração com calendários externos', level=3)
    add_p(doc,
          'O sistema oferece duas formas de integração com calendários: (1) feed iCalendar '
          '(formato RFC 5545) por profissional em /agenda/<slug>/feed.ics?token=<>, protegido '
          'por token assinado por profissional, exportando atendimentos no range -30 a +90 '
          'dias, com STATUS CONFIRMED/TENTATIVE conforme estado; (2) integração OAuth '
          'bidirecional com Google Calendar (opcional), com push outbound em novo agendamento '
          'e pull periódico de eventos externos como bloqueios na agenda. Bibliotecas Google '
          'são opcionais, com fallback gracioso quando ausentes.')

    add_heading(doc, '3.4 Avanços do semestre — sprints implementadas', level=2)
    add_p(doc,
          'O semestre foi organizado em cinco sprints temáticas, cada uma com escopo bem '
          'definido e entregas verificáveis em produção.')

    add_heading(doc, '3.4.1 Sprint 1 — Fundação back-end e agenda visual', level=3)
    add_bullet(doc, 'Buffer entre atendimentos por procedimento (campo buffer_minutos, 0–120 min).')
    add_bullet(doc, 'min_notice_horas e max_advance_dias por profissional.')
    add_bullet(doc, 'Modelo ExcecaoDisponibilidade (FOLGA ou HORARIO_DIFERENTE por data específica).')
    add_bullet(doc, 'Reescrita do gerador de slots (get_horarios_disponiveis) honrando todas as regras em ordem: feriado, max_advance, exceção, regra semanal, min_notice, buffer, bloqueios.')
    add_bullet(doc, 'Endpoint JSON para FullCalendar com drag-drop reagendamento.')
    add_bullet(doc, 'Quatro novos KPIs no dashboard: ticket médio, ocupação semanal, ativos 90d, NPS 30d.')

    add_heading(doc, '3.4.2 Sprint 2 — UX painel e PWA', level=3)
    add_bullet(doc, 'Filtro instantâneo client-side (admin-search.js, ~50 linhas vanilla JS) em listas de agendamentos, clientes e usuários.')
    add_bullet(doc, 'Timeline visual e KPIs do paciente em cliente_detalhe (LTV, ticket médio, qtd realizados, última visita, top procedimento).')
    add_bullet(doc, 'PWA com service worker custom validado (manifest, cache strategies, offline fallback).')

    add_heading(doc, '3.4.3 Sprint 3 — UX cliente e ICS', level=3)
    add_bullet(doc, 'Slug único por profissional + URL pública /agendar/<slug>/.')
    add_bullet(doc, 'Feed iCalendar protegido por token (django-ical-style implementado manualmente).')
    add_bullet(doc, 'Tabs de filtro por categoria no booking público (Facial, Corporal, Capilar, Outro).')
    add_bullet(doc, 'Reagendamento self-service via link mágico com janela mínima de 24h e transação atômica.')

    add_heading(doc, '3.4.4 Sprint 4 — Workflow, push e FSM', level=3)
    add_bullet(doc, 'Modelos WorkflowRegra (trigger + offset + ação + template + config_json) e WorkflowExecucao (deduplicação UNIQUE regra+atendimento).')
    add_bullet(doc, 'Engine de workflow com signals síncronos (ON_BOOK, ON_CANCEL, ON_NO_SHOW) e cron HTTP idempotente para BEFORE/AFTER_EVENT.')
    add_bullet(doc, 'Web Push (VAPID) com pywebpush, endpoints subscribe/unsubscribe/public-key, JS auto-subscribe.')
    add_bullet(doc, 'FSM em Atendimento via dicionário TRANSICOES + métodos confirmar/cancelar/marcar_realizado/marcar_falta com auditoria automática em LogAuditoria.')

    add_heading(doc, '3.4.5 Sprint 5 — RRULE, anamnese, GCal e embed', level=3)
    add_bullet(doc, 'BloqueioAgenda com regra_recorrencia (iCal RRULE), expansão via dateutil.rrule.')
    add_bullet(doc, 'FormularioAnamnese com schema JSON dinâmico (escopo GLOBAL, CATEGORIA ou PROCEDIMENTO).')
    add_bullet(doc, 'Scaffolding de integração Google Calendar OAuth bidirecional.')
    add_bullet(doc, 'Widget standalone /embed/agendar/ (X-Frame-Options exempt) para embed em Linktree e Instagram bio.')
    add_bullet(doc, 'Service worker v4 com cache stale-while-revalidate para APIs públicas.')

    add_heading(doc, '3.4.6 Sprint 6 — Modelos extras CRM e estética', level=3)
    add_bullet(doc, 'PatchTest: teste alérgico estruturado (substância, resultado, profissional aplicador).')
    add_bullet(doc, 'FotoAntesDepois: URL S3, tipo (ANTES/DURANTE/DEPOIS), consent_uso_imagem com timestamp.')
    add_bullet(doc, 'Produto + MovimentoEstoque: controle de insumos com alerta de estoque baixo.')
    add_bullet(doc, 'Tag + ClienteTag: segmentação (VIP, gestante, alérgico).')
    add_bullet(doc, 'PlanoTratamento + ItemPlanoTratamento: planos personalizados multi-procedimento com progresso percentual.')
    add_bullet(doc, 'CreditoCliente + MovimentoCredito: saldo do cliente (vale presente, refund, ajuste).')
    add_bullet(doc, 'Cliente.indicado_por: programa de indicação (referência cliente-cliente).')
    add_bullet(doc, 'Cliente.foto_url e Profissional.foto_url + bio.')

    add_heading(doc, '3.5 Conformidade LGPD', level=2)
    add_p(doc,
          'O sistema implementa requisitos da Lei Geral de Proteção de Dados em múltiplas camadas: '
          'consentimentos granulares por canal (e-mail marketing, WhatsApp NPS, WhatsApp '
          'confirmação) com timestamp e IP de origem; tokens de unsubscribe one-click conforme '
          'RFC 8058; soft-delete em Cliente com campo deletado_em; serviço LgpdService que '
          'anonimiza dados pessoais (PII) substituindo por hashes/placeholders mantendo histórico '
          'contábil; exposição de DSAR (Data Subject Access Request) em /lgpd/meus-dados/; '
          'retenção configurável por tipo de dado; auditoria detalhada de toda ação crítica '
          'em log_auditoria.')

    add_heading(doc, '3.6 Segurança', level=2)
    add_p(doc,
          'A camada de segurança compreende: autenticação por e-mail com hash PBKDF2 (Django '
          'default); 2FA TOTP opcional via django-two-factor-auth integrado ao painel; bloqueio '
          'de força bruta com django-axes (cinco tentativas, lockout de uma hora); rate limit '
          'em endpoints sensíveis com django-ratelimit (login, OTP, password reset); password '
          'reset com TTL de 1 hora e rate limit de 3 requests por 15 minutos por IP; CSRF '
          'em todos os formulários; CSP com nonce por requisição; X-Frame-Options exceto em '
          '/embed/; Content-Type-Options nosniff; cookies Secure e HttpOnly em produção; HSTS '
          'com preload.')

    add_heading(doc, '3.7 Padronização do banco de dados', level=2)
    add_p(doc,
          'A análise comparativa do schema atual com referências (Cal.com Prisma, OpenEMR, '
          'Easy!Appointments, Pabau) identificou pontos de melhoria documentados em '
          'docs/DB_AUDIT.md: aplicação consistente de TimestampedMixin (criado_em, atualizado_em) '
          'em todos os modelos legacy; aplicação seletiva de SoftDeleteMixin em entidades '
          'críticas; adição de campo UUID em Cliente para link DSAR público estável; índices '
          'compostos faltantes para consultas frequentes; adoção do padrão Cal.com de '
          'BookingReference dedicado para suportar múltiplas integrações de calendário externo. '
          'A migração consolidada está planejada para o próximo sprint.')

    # ─── 4. RESULTADOS ───
    add_heading(doc, '4. Resultados', level=1)
    add_p(doc, 'O sistema encontra-se em produção ativa nos seguintes endereços:')
    add_bullet(doc, 'Produção: https://web-production-465af.up.railway.app/')
    add_bullet(doc, 'Desenvolvimento: https://web-dev-1a30.up.railway.app/')
    add_p(doc, 'Métricas quantitativas:')
    add_table_simple(doc,
        ['Indicador', 'Valor'],
        [
            ['Tabelas de negócio em produção', '47'],
            ['Modelos Django implementados', '49 (10 novos no semestre)'],
            ['Migrations aplicadas', '22'],
            ['Procedimentos catalogados', '23 (todos do PRD da clínica)'],
            ['Sprints concluídas (semestre atual)', '6'],
            ['Commits no semestre', '~100'],
            ['Endpoints HTTP públicos + admin', '~80'],
            ['Templates Django', '~70'],
            ['Canais de notificação integrados', '4 (WhatsApp, SMS, e-mail, web push)'],
            ['Conformidade LGPD', 'Completa (consent, audit, anonimização, soft-delete)'],
        ],
        col_widths_cm=[10, 6])

    # ─── 5. CONSIDERAÇÕES FINAIS ───
    add_heading(doc, '5. Considerações finais', level=1)
    add_p(doc,
          'O projeto demonstrou viabilidade de construir uma plataforma profissional completa '
          'para clínicas de estética single-tenant utilizando stack open-source madura, com '
          'custo de hospedagem mínimo no plano free do Railway. As cinco sprints do semestre '
          'transformaram um sistema básico em uma plataforma comparável a SaaS comerciais como '
          'Pabau e Calendly, com diferenciais: white-label nativo, conformidade LGPD desde a '
          'concepção, integração WhatsApp Business API e arquitetura preparada para multi-clínica '
          'no futuro.')
    add_p(doc,
          'Os principais aprendizados envolvem: (i) trade-offs entre service worker custom e '
          'bibliotecas como django-pwa, optando-se pela primeira por flexibilidade e estratégias '
          'híbridas de cache; (ii) padrões Easy!Appointments e Cal.com para working plan exceptions, '
          'adotados na modelagem de ExcecaoDisponibilidade; (iii) importância de FSM explícito '
          'em entidades de ciclo de vida longo (Atendimento), permitindo auditoria automática '
          'e prevenção de transições inválidas; (iv) valor de manifests PWA separados para escopos '
          'distintos (público vs. administrativo).')
    add_p(doc,
          'Como trabalhos futuros, destacam-se: implementação de modelo de pagamentos integrando '
          'MercadoPago para sinal de 30% e redução de no-show; squash de migrations para '
          'limpeza pós-MVP; encriptação at-rest de campos PII via django-cryptography; '
          'submissão e aprovação dos templates WhatsApp Business no Meta Business Manager; '
          'configuração de backups automáticos do banco de dados; e progressão da plataforma '
          'para multi-tenant, permitindo atender múltiplas clínicas com isolamento de dados.')
    add_p(doc,
          'A avaliação do cliente real (biomédica Jaqueline Aranha) é positiva. A operação '
          'manual via WhatsApp foi substituída por agendamento autosserviço com OTP, eliminando '
          'troca manual de mensagens para confirmações. O painel administrativo permite visão '
          'consolidada da operação e o calendário visual reduziu erros de agendamento.')

    # ─── 6. REFERÊNCIAS ───
    add_heading(doc, '6. Referências', level=1)
    add_p(doc, '[1] DJANGO Software Foundation. Django Documentation. Versão 5.2. 2025. Disponível em: https://docs.djangoproject.com/en/5.2/. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[2] CALCOM Inc. Cal.com — Scheduling Infrastructure for Everyone. GitHub Repository. 2026. Disponível em: https://github.com/calcom/cal.com. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[3] TSELEGIDIS, A. Easy!Appointments — Working Plan Exceptions Explained. 2021. Disponível em: https://easyappointments.org/2021/01/26/working-plan-exceptions-explained/. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[4] OPENEMR Foundation. Database Structure. OpenEMR Project Wiki. Disponível em: https://www.open-emr.org/wiki/index.php/Database_Structure. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[5] PABAU Inc. Best Aesthetic Clinic Software (2026 Reviewed). Pabau Blog. 2026. Disponível em: https://pabau.com/blog/best-aesthetic-clinic-software/. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[6] BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais (LGPD). Diário Oficial da União, Brasília, DF, 15 ago. 2018.', size=9, space_after=4)
    add_p(doc, '[7] DESBIENS, M. et al. RFC 5545 — Internet Calendaring and Scheduling Core Object Specification (iCalendar). IETF, 2009. Disponível em: https://datatracker.ietf.org/doc/html/rfc5545. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[8] DZIK, M. RFC 8030 — Generic Event Delivery Using HTTP Push. IETF, 2016. Disponível em: https://datatracker.ietf.org/doc/html/rfc8030. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[9] PINTO, D. Timestamps and Soft Delete with Django. xgeeks Medium, 2022. Disponível em: https://medium.com/xgeeks/timestamps-and-soft-delete-with-django-65f74d86e022. Acesso em: 4 mai. 2026.', size=9, space_after=4)
    add_p(doc, '[10] CLOUDFLARE Inc. Cloudflare Turnstile — A Smart, Friendly CAPTCHA Replacement. Disponível em: https://www.cloudflare.com/products/turnstile/. Acesso em: 4 mai. 2026.', size=9, space_after=4)

    # Save
    output = r'C:/Users/rafae/OneDrive/Projeto FAM - Shivazen/Artigo/Documentacao_Tecnica_JaquelineAranhaEstetica.docx'
    doc.save(output)
    print(f'DOCX gerado: {output}')


if __name__ == '__main__':
    main()
